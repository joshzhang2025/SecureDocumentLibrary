from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

import yaml

MAX_BYTES, MAX_CHARS = 20 * 1024 * 1024, 500_000

@dataclass(frozen=True)
class ParsedPart:
    title: str
    text: str
    document_type: str
    part_name: str | None = None

def _text(path: Path) -> str:
    if path.stat().st_size > MAX_BYTES: raise ValueError("File exceeds size limit")
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try: return data.decode(encoding)[:MAX_CHARS]
        except UnicodeDecodeError: pass
    raise ValueError("Text cannot be decoded")

def parse(path: Path) -> list[ParsedPart]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt", ".yaml", ".yml", ".json"}:
        text = _text(path)
        if suffix == ".json": json.loads(text)
        if suffix in {".yaml", ".yml"}: yaml.safe_load(text)
        headings = [line[2:].strip() for line in text.splitlines() if line.startswith("# ")]
        return [ParsedPart(headings[0] if headings else path.stem, text, "markdown" if suffix == ".md" else "text")]
    if suffix == ".docx":
        from docx import Document
        document = Document(path); lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        lines.extend(" | ".join(cell.text.replace("\n", " ").strip() for cell in row.cells) for table in document.tables for row in table.rows)
        return [ParsedPart(document.core_properties.title or path.stem, "\n".join(lines)[:MAX_CHARS], "docx")]
    if suffix == ".xlsx":
        from openpyxl import load_workbook
        book = load_workbook(path, read_only=True, data_only=True, keep_links=False); parts = []
        try:
            for sheet in list(book.worksheets)[:50]:
                rows = []
                for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    if row_number > 100_000: break
                    values = ["" if value is None else (value.isoformat() if isinstance(value, (date, datetime)) else str(value)) for value in row[:500]]
                    if any(values): rows.append(" | ".join(values))
                    if sum(len(value) for value in rows) > MAX_CHARS: break
                parts.append(ParsedPart(f"{path.stem} - {sheet.title}", "\n".join(rows)[:MAX_CHARS], "xlsx_sheet", sheet.title))
        finally: book.close()
        return parts
    raise ValueError(f"Unsupported file extension: {suffix}")

